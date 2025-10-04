"""
Document processor service
Handles different document types (PDF, DOC/DOCX, Images) with appropriate extraction methods
"""

import io
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document

from utils.document_detector import detect_document_type
from services.ocr_service import VisionService
from services.classification_service import ClassificationService

class DocumentProcessor:
    """
    Unified document processing service that handles multiple document types
    """
    
    def __init__(self):
        """Initialize with required services"""
        print("[DOC-PROCESSOR] Initializing document processor service...")
        self.vision_service = VisionService()
        self.classification_service = ClassificationService()
        print("[DOC-PROCESSOR] ✅ Document processor service ready")
        
    async def process_document(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Process document based on its detected type
        
        Args:
            file_bytes: Raw bytes of the document file
            filename: Original filename with extension
            
        Returns:
            Dict with processing results including document type, extracted text, and classification
        """
        # Detect document type
        doc_type = detect_document_type(filename)
        print(f"[DOC-PROCESSOR] Detected document type: {doc_type} for file {filename}")
        
        # Extract text based on document type
        if doc_type == "pdf":
            text = self._extract_pdf_text(file_bytes)
            print(f"[DOC-PROCESSOR] Extracted {len(text)} characters from PDF")
        elif doc_type in ["doc", "docx"]:
            text = self._extract_doc_text(file_bytes, doc_type)
            print(f"[DOC-PROCESSOR] Extracted {len(text)} characters from {doc_type.upper()}")
        elif doc_type == "image":
            # Use existing OCR pipeline
            ocr_result = self.vision_service.extract_text(file_bytes)
            text = ocr_result.text
            print(f"[DOC-PROCESSOR] Extracted {len(text)} characters from image using OCR")
        else:
            print(f"[DOC-PROCESSOR] ❌ Unsupported document type: {doc_type}")
            return {
                "success": False,
                "error": f"Unsupported document type: {doc_type}",
                "doc_type": doc_type
            }
        
        if not text or len(text.strip()) == 0:
            print(f"[DOC-PROCESSOR] ⚠️ No text extracted from {doc_type} document")
            return {
                "success": False,
                "error": f"Could not extract text from {doc_type} document",
                "doc_type": doc_type,
                "text": "",
                "text_length": 0
            }
        
        # Classify the document using existing classification service
        try:
            print(f"[DOC-PROCESSOR] Classifying document content...")
            classification = await self.classification_service.classify_document(text)
            print(f"[DOC-PROCESSOR] ✅ Document classified successfully")
        except Exception as e:
            print(f"[DOC-PROCESSOR] ❌ Classification error: {str(e)}")
            classification = {"error": f"Classification failed: {str(e)}"}
        
        # Return unified response format
        return {
            "success": True,
            "doc_type": doc_type,
            "text": text[:1000] + "..." if len(text) > 1000 else text,  # Truncate preview for response
            "text_length": len(text),
            "classification": classification
        }
    
    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_bytes: Raw bytes of PDF file
            
        Returns:
            Extracted text string
        """
        pdf_text = ""
        try:
            with io.BytesIO(file_bytes) as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                print(f"[DOC-PROCESSOR] PDF has {len(pdf_reader.pages)} pages")
                
                for page_num in range(len(pdf_reader.pages)):
                    page_text = pdf_reader.pages[page_num].extract_text() or ""
                    pdf_text += page_text + "\n\n"
                    
            return pdf_text
        except Exception as e:
            error_msg = f"PDF extraction error: {str(e)}"
            print(f"[DOC-PROCESSOR] ❌ {error_msg}")
            return f"Error extracting PDF text: {str(e)}"
    
    def _extract_doc_text(self, file_bytes: bytes, doc_type: str) -> str:
        """
        Extract text from DOC/DOCX file
        
        Args:
            file_bytes: Raw bytes of DOC/DOCX file
            doc_type: Either "doc" or "docx"
            
        Returns:
            Extracted text string
        """
        try:
            doc_text = ""
            with io.BytesIO(file_bytes) as doc_file:
                doc = Document(doc_file)
                print(f"[DOC-PROCESSOR] Document has {len(doc.paragraphs)} paragraphs")
                
                for para in doc.paragraphs:
                    doc_text += para.text + "\n"
                    
                # Extract text from tables if present
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            doc_text += cell.text + " | "
                        doc_text += "\n"
                    doc_text += "\n"
                    
            return doc_text
        except Exception as e:
            error_msg = f"{doc_type.upper()} extraction error: {str(e)}"
            print(f"[DOC-PROCESSOR] ❌ {error_msg}")
            return f"Error extracting {doc_type.upper()} text: {str(e)}"